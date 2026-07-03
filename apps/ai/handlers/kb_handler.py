"""
KB processing pipeline:
  URL  → httpx fetch → BeautifulSoup text extraction
  File → download presigned S3 URL → parse by type → chunk → embed → upsert Pinecone
"""

import os
import io
import copy
import ipaddress
import inspect
import json
import secrets
import socket
import time
from datetime import datetime, timezone
from threading import RLock
from typing import Optional
from urllib.parse import urlparse, urlunparse

from utils.logger import logger, redact_sensitive
from utils.metrics import emit_metric
from utils.pinecone_client import pinecone_client

# ── lazy imports (heavy deps loaded once) ────────────────────────────────────

def _pinecone():
    return pinecone_client()

def _index():
    pc = _pinecone()
    return pc.Index(os.environ.get("PINECONE_INDEX", "quickvoice-kb"))

EMBEDDING_MODEL = os.environ.get("PINECONE_EMBEDDING_MODEL", "llama-text-embed-v2")
EMBEDDING_TRUNCATE = os.environ.get("PINECONE_EMBEDDING_TRUNCATE", "END")
CHUNK_SIZE = 500       # characters (not tokens — keeps it dependency-light)
CHUNK_OVERLAP = 50
MAX_DOWNLOAD_BYTES = int(os.environ.get("KB_MAX_DOWNLOAD_BYTES", str(10 * 1024 * 1024)))
MAX_CHUNKS_PER_DOCUMENT = int(os.environ.get("KB_MAX_CHUNKS_PER_DOCUMENT", "500"))
MAX_DOCUMENTS_PER_JOB = int(os.environ.get("KB_MAX_DOCUMENTS_PER_JOB", "50"))
ALLOWED_SCHEMES = {"http", "https"}
ALLOWED_HOSTS = [
    host.strip().lower()
    for host in os.environ.get("KB_ALLOWED_HOSTS", "").split(",")
    if host.strip()
]
ALLOWED_CONTENT_TYPES = (
    "text/",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument",
    "application/vnd.ms-excel",
    "application/json",
    "application/csv",
    "application/octet-stream",
)

TERMINAL_JOB_STATUSES = {"succeeded", "partial_failed", "failed", "canceled"}
PROCESSED_DOCUMENT_STATUSES = {"ok", "error", "canceled"}
_KB_JOBS: dict[str, dict] = {}
_KB_JOB_LOCK = RLock()


class KbJobValidationError(ValueError):
    def __init__(
        self,
        *,
        status_code: int,
        code: str,
        user_message: str,
        retryable: bool = False,
        details: dict | None = None,
    ):
        super().__init__(user_message)
        self.status_code = status_code
        self.code = code
        self.user_message = user_message
        self.retryable = retryable
        self.details = details or {}

    def to_detail(self) -> dict:
        return {
            "code": self.code,
            "userMessage": self.user_message,
            "retryable": self.retryable,
            "details": self.details,
        }


# ── KB job state ─────────────────────────────────────────────────────────────

def create_kb_job(payload: dict) -> dict:
    normalized = _validate_process_payload(payload)
    job_id = _new_job_id()
    now = _now_iso()
    job = {
        "jobId": job_id,
        "agentId": normalized["agentId"],
        "organizationId": normalized["organizationId"],
        "status": "queued",
        "stage": "queued",
        "progress": _progress_for_documents([]),
        "documents": [_queued_document_state(doc) for doc in normalized["documents"]],
        "statusUrl": f"/kb/jobs/{job_id}",
        "createdAt": now,
        "updatedAt": now,
        "finishedAt": None,
        "error": None,
        "_payload": copy.deepcopy(normalized),
        "_cancelRequested": False,
    }
    job["progress"] = _progress_for_documents(job["documents"])

    with _KB_JOB_LOCK:
        _KB_JOBS[job_id] = job

    emit_metric(
        "kb_job_created",
        agent_id=normalized["agentId"],
        organization_id=normalized["organizationId"],
        job_id=job_id,
        documents=len(normalized["documents"]),
    )
    logger.info(
        "[kb] job created {}",
        redact_sensitive({"jobId": job_id, "agentId": normalized["agentId"], "documents": len(normalized["documents"])}),
    )
    return _public_job(job)


def get_kb_job(job_id: str) -> dict:
    with _KB_JOB_LOCK:
        job = _KB_JOBS.get(job_id)
        if not job:
            raise KeyError(job_id)
        return _public_job(job)


def cancel_kb_job(job_id: str) -> dict:
    with _KB_JOB_LOCK:
        job = _KB_JOBS.get(job_id)
        if not job:
            raise KeyError(job_id)
        if job["status"] in TERMINAL_JOB_STATUSES:
            return _public_job(job)

        job["_cancelRequested"] = True
        if job["status"] == "queued":
            _mark_remaining_documents_canceled_locked(job)
            _finalize_job_locked(job, canceled=True)
        else:
            job["status"] = "canceling"
            job["stage"] = "canceling"
            job["updatedAt"] = _now_iso()
        return _public_job(job)


def retry_kb_job(job_id: str) -> dict:
    with _KB_JOB_LOCK:
        job = _KB_JOBS.get(job_id)
        if not job:
            raise KeyError(job_id)
        failed_ids = {doc["kbId"] for doc in job["documents"] if doc.get("status") == "error"}
        if not failed_ids:
            raise KbJobValidationError(
                status_code=400,
                code="KB_JOB_HAS_NO_FAILED_DOCUMENTS",
                user_message="There are no failed knowledge sources to retry for this job.",
                retryable=False,
            )
        payload = copy.deepcopy(job["_payload"])

    payload["documents"] = [doc for doc in payload["documents"] if doc.get("kbId") in failed_ids]
    return create_kb_job(payload)


async def run_kb_job(job_id: str) -> dict:
    with _KB_JOB_LOCK:
        job = _KB_JOBS.get(job_id)
        if not job:
            raise KeyError(job_id)
        if job["status"] in TERMINAL_JOB_STATUSES:
            return _public_job(job)
        if job.get("_cancelRequested"):
            _mark_remaining_documents_canceled_locked(job)
            _finalize_job_locked(job, canceled=True)
            return _public_job(job)

        job["status"] = "running"
        job["stage"] = "processing"
        job["updatedAt"] = _now_iso()
        payload = copy.deepcopy(job["_payload"])

    async def progress(event: dict) -> None:
        _apply_job_progress(job_id, event)

    def should_cancel() -> bool:
        with _KB_JOB_LOCK:
            current = _KB_JOBS.get(job_id)
            return bool(current and current.get("_cancelRequested"))

    try:
        results = await process_documents(payload, progress=progress, should_cancel=should_cancel)
    except Exception as exc:
        detail = _job_error_detail(exc)
        with _KB_JOB_LOCK:
            job = _KB_JOBS[job_id]
            job["status"] = "failed"
            job["stage"] = "failed"
            job["error"] = detail
            job["updatedAt"] = _now_iso()
            job["finishedAt"] = job["updatedAt"]
            for doc in job["documents"]:
                if doc.get("status") not in PROCESSED_DOCUMENT_STATUSES:
                    doc.update(_document_error_fields(exc, budget=_budget_for_agent(job["agentId"], job["_payload"])))
            job["progress"] = _progress_for_documents(job["documents"])
            public = _public_job(job)
        emit_metric("kb_job_completed", status="failed", job_id=job_id, error_code=detail["code"])
        logger.error("[kb] job failed {}", redact_sensitive({"jobId": job_id, "code": detail["code"]}))
        return public

    with _KB_JOB_LOCK:
        job = _KB_JOBS[job_id]
        for result in results:
            _apply_document_result_locked(job, result)
        if job.get("_cancelRequested") or any(result.get("status") == "canceled" for result in results):
            _mark_remaining_documents_canceled_locked(job)
            _finalize_job_locked(job, canceled=True)
        else:
            _finalize_job_locked(job)
        public = _public_job(job)

    emit_metric(
        "kb_job_completed",
        status=public["status"],
        job_id=job_id,
        agent_id=public["agentId"],
        documents=public["progress"]["total"],
        succeeded=public["progress"]["succeeded"],
        failed=public["progress"]["failed"],
    )
    logger.info("[kb] job completed {}", redact_sensitive({"jobId": job_id, "status": public["status"]}))
    return public


def _new_job_id() -> str:
    return f"kbjob_{secrets.token_urlsafe(12)}"


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _validate_process_payload(payload: dict) -> dict:
    if not isinstance(payload, dict):
        raise KbJobValidationError(
            status_code=400,
            code="KB_REQUEST_INVALID",
            user_message="Knowledge processing requires a valid request body.",
            retryable=False,
        )

    agent_id = str(payload.get("agentId") or "").strip()
    if not agent_id:
        raise KbJobValidationError(
            status_code=400,
            code="KB_AGENT_ID_REQUIRED",
            user_message="Choose an agent before processing knowledge sources.",
            retryable=False,
        )

    organization_id = str(payload.get("organizationId") or "").strip()
    if not organization_id:
        raise KbJobValidationError(
            status_code=400,
            code="KB_ORGANIZATION_ID_REQUIRED",
            user_message="Choose an organization before processing knowledge sources.",
            retryable=False,
        )

    documents = payload.get("documents")
    if not isinstance(documents, list) or not documents:
        raise KbJobValidationError(
            status_code=400,
            code="KB_DOCUMENTS_REQUIRED",
            user_message="Add at least one knowledge source before starting processing.",
            retryable=False,
        )
    if any(not isinstance(doc, dict) for doc in documents):
        raise KbJobValidationError(
            status_code=400,
            code="KB_DOCUMENT_INVALID",
            user_message="Each knowledge source must include document metadata.",
            retryable=False,
        )

    budget = _budget_for_agent(agent_id, payload)
    max_documents = budget["max_documents_per_job"]
    if len(documents) > max_documents:
        raise KbJobValidationError(
            status_code=413,
            code="KB_DOCUMENT_LIMIT_EXCEEDED",
            user_message=f"This job has {len(documents)} knowledge sources, but the limit is {max_documents}.",
            retryable=False,
            details={"maxDocumentsPerJob": max_documents, "documentCount": len(documents)},
        )

    normalized = copy.deepcopy(payload)
    normalized["agentId"] = agent_id
    normalized["organizationId"] = organization_id
    normalized["documents"] = [copy.deepcopy(doc) for doc in documents]
    return normalized


def _queued_document_state(doc: dict) -> dict:
    kb_id = str(doc.get("kbId") or "")
    return {
        "kbId": kb_id,
        "name": doc.get("name") or doc.get("originalFileName") or kb_id,
        "sourceType": doc.get("sourceType", "TXT"),
        "status": "queued",
        "stage": "queued",
    }


def _public_job(job: dict) -> dict:
    public = {
        key: copy.deepcopy(value)
        for key, value in job.items()
        if not key.startswith("_") and value is not None
    }
    public["documents"] = [_public_document(doc) for doc in job.get("documents", [])]
    return public


def _public_document(doc: dict) -> dict:
    return {
        key: copy.deepcopy(value)
        for key, value in doc.items()
        if not key.startswith("_") and value is not None
    }


def _progress_for_documents(documents: list[dict]) -> dict:
    total = len(documents)
    succeeded = sum(1 for doc in documents if doc.get("status") == "ok")
    failed = sum(1 for doc in documents if doc.get("status") == "error")
    canceled = sum(1 for doc in documents if doc.get("status") == "canceled")
    processed = succeeded + failed + canceled
    percent = int((processed / total) * 100) if total else 0
    return {
        "total": total,
        "processed": processed,
        "succeeded": succeeded,
        "failed": failed,
        "canceled": canceled,
        "percent": percent,
    }


def _find_job_document(job: dict, kb_id: str) -> dict | None:
    for doc in job.get("documents", []):
        if doc.get("kbId") == kb_id:
            return doc
    return None


def _apply_job_progress(job_id: str, event: dict) -> None:
    with _KB_JOB_LOCK:
        job = _KB_JOBS.get(job_id)
        if not job:
            return
        if job["status"] != "canceling":
            job["status"] = "running"
            job["stage"] = "processing"
        _apply_document_result_locked(job, event)
        job["updatedAt"] = _now_iso()


def _apply_document_result_locked(job: dict, result: dict) -> None:
    kb_id = result.get("kbId")
    if not kb_id:
        return
    doc = _find_job_document(job, kb_id)
    if not doc:
        doc = {"kbId": kb_id, "name": result.get("name") or kb_id, "sourceType": result.get("sourceType", "TXT")}
        job["documents"].append(doc)

    for key in ("name", "sourceType", "status", "stage", "chunks", "code", "userMessage", "retryable", "details", "error"):
        if key in result:
            doc[key] = copy.deepcopy(result[key])

    if doc.get("status") == "ok":
        doc.setdefault("stage", "indexed")
        for key in ("code", "userMessage", "retryable", "details", "error"):
            doc.pop(key, None)
    elif doc.get("status") == "error":
        doc.setdefault("stage", "failed")
    elif doc.get("status") == "canceled":
        doc.setdefault("stage", "canceled")

    job["progress"] = _progress_for_documents(job["documents"])
    job["updatedAt"] = _now_iso()


def _mark_remaining_documents_canceled_locked(job: dict) -> None:
    for doc in job.get("documents", []):
        if doc.get("status") not in PROCESSED_DOCUMENT_STATUSES:
            doc.update(_canceled_document_fields())
    job["progress"] = _progress_for_documents(job["documents"])
    job["updatedAt"] = _now_iso()


def _finalize_job_locked(job: dict, *, canceled: bool = False) -> None:
    job["progress"] = _progress_for_documents(job["documents"])
    progress = job["progress"]
    if canceled:
        job["status"] = "canceled"
        job["stage"] = "canceled"
    elif progress["failed"] == 0 and progress["canceled"] == 0:
        job["status"] = "succeeded"
        job["stage"] = "completed"
    elif progress["succeeded"] == 0 and progress["canceled"] == 0:
        job["status"] = "failed"
        job["stage"] = "failed"
    else:
        job["status"] = "partial_failed"
        job["stage"] = "completed"
    job["updatedAt"] = _now_iso()
    job["finishedAt"] = job["updatedAt"]


def _canceled_document_fields() -> dict:
    user_message = "Knowledge processing was canceled before this source finished."
    return {
        "status": "canceled",
        "stage": "canceled",
        "code": "KB_JOB_CANCELED",
        "userMessage": user_message,
        "retryable": True,
        "error": user_message,
    }


def _job_error_detail(exc: Exception) -> dict:
    if isinstance(exc, KbJobValidationError):
        return exc.to_detail()
    error = _document_error_fields(exc, budget={})
    return {
        "code": error["code"],
        "userMessage": error["userMessage"],
        "retryable": error["retryable"],
        "details": error.get("details", {}),
    }


def _document_error_fields(exc: Exception, *, budget: dict) -> dict:
    message = str(exc)
    details: dict = {}
    retryable = False

    if "url is required for URL source type" in message:
        code = "KB_URL_REQUIRED"
        user_message = "Add a URL for this knowledge source and try again."
    elif "presignedUrl is required for file source types" in message:
        code = "KB_FILE_URL_REQUIRED"
        user_message = "Upload the file again so QuickVoice can read it."
        retryable = True
    elif "Only http and https URLs are allowed" in message:
        code = "KB_URL_UNSUPPORTED_SCHEME"
        user_message = "Use an http or https URL for this knowledge source."
    elif "URL host is required" in message:
        code = "KB_URL_HOST_REQUIRED"
        user_message = "Enter a complete URL with a host name."
    elif "URL credentials are not allowed" in message:
        code = "KB_URL_CREDENTIALS_NOT_ALLOWED"
        user_message = "Remove credentials from the URL before processing this source."
    elif "Local hosts are not allowed" in message or "Private, local, or reserved IP addresses" in message:
        code = "KB_URL_PRIVATE_HOST"
        user_message = "Use a public URL for this knowledge source."
    elif "URL host is not allowed" in message:
        code = "KB_URL_HOST_NOT_ALLOWED"
        user_message = "This URL host is not allowed for knowledge ingestion."
    elif "URL host could not be resolved" in message:
        code = "KB_URL_HOST_UNRESOLVED"
        user_message = "QuickVoice could not reach this URL host. Check the URL and try again."
        retryable = True
    elif "download exceeds" in message:
        code = "KB_DOWNLOAD_TOO_LARGE"
        user_message = "This knowledge source is larger than the allowed download size."
    elif "Too many redirects" in message:
        code = "KB_TOO_MANY_REDIRECTS"
        user_message = "This URL redirects too many times. Use the final destination URL and try again."
    elif "Unsupported URL content type" in message:
        code = "KB_UNSUPPORTED_URL_CONTENT_TYPE"
        user_message = "This URL does not point to a supported web page."
    elif "Unsupported file content type" in message:
        code = "KB_UNSUPPORTED_FILE_CONTENT_TYPE"
        user_message = "This file type is not supported for knowledge ingestion."
    elif "Extracted text is empty" in message or "No chunks produced" in message:
        code = "KB_EMPTY_TEXT"
        user_message = "No readable text was found in this knowledge source."
    elif "Document exceeds chunk budget" in message:
        max_chunks = budget.get("max_chunks_per_document") or MAX_CHUNKS_PER_DOCUMENT
        code = "KB_CHUNK_LIMIT_EXCEEDED"
        user_message = f"This knowledge source is too large to index. Reduce it to {max_chunks} chunks or fewer and try again."
        details = {"maxChunksPerDocument": max_chunks}
    elif "looks like a Google API key" in message or ("Invalid API key" in message and ("401" in message or "Unauthorized" in message)):
        code = "KB_VECTOR_STORE_API_KEY_INVALID"
        user_message = "PINECONE_API_KEY was rejected or is not a Pinecone key. Check the AI service environment uses a valid Pinecone API key for this project."
    elif "PINECONE_API_KEY" in message:
        code = "KB_VECTOR_STORE_API_KEY_MISSING"
        user_message = "Knowledge processing requires PINECONE_API_KEY in the AI service environment."
    else:
        code = "KB_PROCESSING_FAILED"
        user_message = "QuickVoice could not process this knowledge source. Try again later."
        retryable = True

    return {
        "status": "error",
        "stage": "failed",
        "code": code,
        "userMessage": user_message,
        "retryable": retryable,
        "details": details,
        "error": user_message,
    }


async def _notify_progress(progress, event: dict) -> None:
    if not progress:
        return
    result = progress(event)
    if inspect.isawaitable(result):
        await result


# ── text extraction ──────────────────────────────────────────────────────────

async def fetch_url(url: str) -> str:
    """Download a web page and extract visible text via BeautifulSoup."""
    import httpx  # type: ignore
    from bs4 import BeautifulSoup  # type: ignore

    content = await download_bytes(url, max_bytes=MAX_DOWNLOAD_BYTES, expected_content="html")
    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


async def download_bytes(
    url: str,
    *,
    max_bytes: int = MAX_DOWNLOAD_BYTES,
    expected_content: str | None = None,
) -> bytes:
    import httpx  # type: ignore

    safe_url = validate_ingest_url(url)
    async with httpx.AsyncClient(follow_redirects=False, timeout=60) as client:
        current_url = safe_url
        for _ in range(5):
            async with client.stream("GET", current_url, headers={"User-Agent": "QuickVoice-KB/1.0"}) as resp:
                if resp.status_code in {301, 302, 303, 307, 308}:
                    redirect_url = resp.headers.get("location")
                    if not redirect_url:
                        raise ValueError("Redirect response missing Location header")
                    current_url = validate_ingest_url(str(resp.url.join(redirect_url)))
                    continue

                resp.raise_for_status()
                _validate_content_type(resp.headers.get("content-type", ""), expected_content)
                chunks: list[bytes] = []
                total = 0
                async for chunk in resp.aiter_bytes():
                    total += len(chunk)
                    if total > max_bytes:
                        raise ValueError(f"download exceeds {max_bytes} byte limit")
                    chunks.append(chunk)
                return b"".join(chunks)

        raise ValueError("Too many redirects")


def validate_ingest_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme.lower() not in ALLOWED_SCHEMES:
        raise ValueError("Only http and https URLs are allowed for KB ingestion")
    if not parsed.hostname:
        raise ValueError("URL host is required")
    if parsed.username or parsed.password:
        raise ValueError("URL credentials are not allowed")

    host = parsed.hostname.rstrip(".").lower()
    if host == "localhost" or host.endswith(".localhost"):
        raise ValueError("Local hosts are not allowed for KB ingestion")
    if ALLOWED_HOSTS and not any(host == allowed or host.endswith(f".{allowed}") for allowed in ALLOWED_HOSTS):
        raise ValueError("URL host is not allowed for KB ingestion")

    _validate_public_host(host)
    return urlunparse(parsed._replace(scheme=parsed.scheme.lower(), netloc=parsed.netloc.lower()))


def _validate_public_host(host: str) -> None:
    try:
        addresses = [ipaddress.ip_address(host)]
    except ValueError:
        try:
            infos = socket.getaddrinfo(host, None, type=socket.SOCK_STREAM)
        except socket.gaierror as exc:
            raise ValueError("URL host could not be resolved") from exc
        addresses = []
        for info in infos:
            addresses.append(ipaddress.ip_address(info[4][0]))

    for address in addresses:
        if (
            address.is_private
            or address.is_loopback
            or address.is_link_local
            or address.is_reserved
            or address.is_multicast
            or address.is_unspecified
        ):
            raise ValueError("Private, local, or reserved IP addresses are not allowed for KB ingestion")


def _validate_content_type(content_type: str, expected_content: str | None) -> None:
    normalized = (content_type or "").split(";", 1)[0].strip().lower()
    if not normalized:
        return
    if expected_content == "html" and normalized not in {"text/html", "application/xhtml+xml"}:
        raise ValueError(f"Unsupported URL content type: {normalized}")
    if expected_content != "html" and not any(normalized.startswith(prefix) for prefix in ALLOWED_CONTENT_TYPES):
        raise ValueError(f"Unsupported file content type: {normalized}")


def parse_pdf(content: bytes) -> str:
    import fitz  # type: ignore  (pymupdf)
    doc = fitz.open(stream=content, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def parse_docx(content: bytes) -> str:
    from docx import Document  # type: ignore
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def parse_xlsx(content: bytes) -> str:
    import openpyxl  # type: ignore
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_text = "\t".join(str(c) for c in row if c is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def parse_xls(content: bytes) -> str:
    import xlrd  # type: ignore
    wb = xlrd.open_workbook(file_contents=content)
    parts = []
    for sheet in wb.sheets():
        for rx in range(sheet.nrows):
            row_text = "\t".join(str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols))
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def parse_file(content: bytes, source_type: str) -> str:
    st = source_type.upper()
    if st == "PDF":
        return parse_pdf(content)
    if st == "DOCX":
        return parse_docx(content)
    if st == "XLSX":
        return parse_xlsx(content)
    if st == "XLS":
        return parse_xls(content)
    # TXT / CSV and any other plain-text type
    return content.decode("utf-8", errors="replace")


# ── chunking ─────────────────────────────────────────────────────────────────

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Simple character-based sliding window splitter."""
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += size - overlap
    return chunks


# ── embedding ─────────────────────────────────────────────────────────────────

def _embedding_values(response) -> list[list[float]]:
    data = response.get("data", []) if isinstance(response, dict) else getattr(response, "data", [])
    values: list[list[float]] = []
    for item in data:
        vector = item.get("values") if isinstance(item, dict) else getattr(item, "values", None)
        if vector is None:
            raise ValueError("Pinecone embedding response did not include values")
        values.append(list(vector))
    return values


async def embed_chunks(chunks: list[str]) -> list[list[float]]:
    """Batch-embed chunks using Pinecone Inference."""
    import asyncio
    pc = _pinecone()

    BATCH = 100
    all_embeddings: list[list[float]] = []

    for i in range(0, len(chunks), BATCH):
        batch = chunks[i : i + BATCH]
        result = await asyncio.to_thread(
            pc.inference.embed,
            model=EMBEDDING_MODEL,
            inputs=batch,
            parameters={"input_type": "passage", "truncate": EMBEDDING_TRUNCATE},
        )
        all_embeddings.extend(_embedding_values(result))

    return all_embeddings


# ── Pinecone upsert ───────────────────────────────────────────────────────────

def upsert_to_pinecone(
    chunks: list[str],
    embeddings: list[list[float]],
    namespace: str,
    kb_id: str,
    doc_name: str,
) -> None:
    index = _index()
    _delete_kb_vectors(index=index, namespace=namespace, kb_id=kb_id)
    vectors = [
        {
            "id": f"{kb_id}#{i}",
            "values": emb,
            "metadata": {
                "kbId": kb_id,
                "name": doc_name,
                "chunkIdx": i,
                "text": chunk[:1000],  # store truncated text for retrieval
            },
        }
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings))
    ]
    # Upsert in batches of 100
    batch_size = 100
    for start in range(0, len(vectors), batch_size):
        index.upsert(vectors=vectors[start : start + batch_size], namespace=namespace)


def delete_kb_vectors(*, namespace: str, kb_id: str) -> None:
    _delete_kb_vectors(index=_index(), namespace=namespace, kb_id=kb_id)


def _delete_kb_vectors(*, index, namespace: str, kb_id: str) -> None:
    try:
        index.delete(filter={"kbId": {"$eq": kb_id}}, namespace=namespace)
    except Exception as exc:
        if _is_pinecone_namespace_not_found(exc):
            logger.info(
                "[kb] pinecone namespace missing during delete; skipping {}",
                redact_sensitive({"namespace": namespace, "kbId": kb_id}),
            )
            return
        raise


def _is_pinecone_namespace_not_found(exc: Exception) -> bool:
    status = getattr(exc, "status", None) or getattr(exc, "status_code", None)
    message_parts = [str(exc)]
    body = getattr(exc, "body", None)
    if body:
        message_parts.append(str(body))
    message = " ".join(message_parts)
    has_404 = status in {404, "404"} or "(404)" in message
    return has_404 and "Namespace not found" in message


# ── main entry ────────────────────────────────────────────────────────────────

async def process_documents(payload: dict, progress=None, should_cancel=None) -> list[dict]:
    """
    Process each document in the payload. Returns a list of
    structured per-document results — 'ok' on success, 'error' on failure.
    """
    payload = _validate_process_payload(payload)
    agent_id: str = payload["agentId"]
    documents: list[dict] = payload["documents"]
    budget = _budget_for_agent(agent_id, payload)
    results: list[dict] = []

    for doc in documents:
        kb_id: str = doc["kbId"]
        name: str = doc.get("name", kb_id)
        source_type: str = doc.get("sourceType", "TXT")
        url: Optional[str] = doc.get("url")
        presigned_url: Optional[str] = doc.get("presignedUrl")
        result_base = {"kbId": kb_id, "name": name, "sourceType": source_type}

        if should_cancel and should_cancel():
            result = {**result_base, **_canceled_document_fields()}
            await _notify_progress(progress, result)
            results.append(result)
            continue

        started = time.perf_counter()

        try:
            # 1. Extract text
            await _notify_progress(progress, {**result_base, "status": "running", "stage": "extracting"})
            if source_type.upper() == "URL":
                if not url:
                    raise ValueError("url is required for URL source type")
                validate_ingest_url(url)
                text = await fetch_url(url)
            else:
                if not presigned_url:
                    raise ValueError("presignedUrl is required for file source types")
                validate_ingest_url(presigned_url)
                content = await download_bytes(presigned_url)
                text = parse_file(content, source_type)

            if not text.strip():
                raise ValueError("Extracted text is empty")

            # 2. Chunk
            if should_cancel and should_cancel():
                result = {**result_base, **_canceled_document_fields()}
                await _notify_progress(progress, result)
                results.append(result)
                continue
            await _notify_progress(progress, {**result_base, "status": "running", "stage": "chunking"})
            chunks = chunk_text(text)
            if not chunks:
                raise ValueError("No chunks produced")
            if len(chunks) > budget["max_chunks_per_document"]:
                raise ValueError(f"Document exceeds chunk budget of {budget['max_chunks_per_document']}")

            # 3. Embed
            if should_cancel and should_cancel():
                result = {**result_base, **_canceled_document_fields()}
                await _notify_progress(progress, result)
                results.append(result)
                continue
            await _notify_progress(
                progress,
                {**result_base, "status": "running", "stage": "embedding", "chunks": len(chunks)},
            )
            embeddings = await embed_chunks(chunks)

            # 4. Upsert to Pinecone (namespace = agentId for per-agent isolation)
            if should_cancel and should_cancel():
                result = {**result_base, **_canceled_document_fields()}
                await _notify_progress(progress, result)
                results.append(result)
                continue
            await _notify_progress(
                progress,
                {**result_base, "status": "running", "stage": "indexing", "chunks": len(chunks)},
            )
            upsert_to_pinecone(chunks, embeddings, namespace=agent_id, kb_id=kb_id, doc_name=name)

            emit_metric(
                "kb_document_processed",
                status="ok",
                agent_id=agent_id,
                kb_id=kb_id,
                source_type=source_type,
                text_chars=len(text),
                chunks=len(chunks),
                latency_ms=int((time.perf_counter() - started) * 1000),
            )
            logger.info(
                "[kb] processed {}",
                redact_sensitive({"kbId": kb_id, "chunks": len(chunks), "sourceType": source_type}),
            )
            result = {**result_base, "status": "ok", "stage": "indexed", "chunks": len(chunks)}
            await _notify_progress(progress, result)
            results.append(result)

        except Exception as exc:
            error = _document_error_fields(exc, budget=budget)
            emit_metric(
                "kb_document_processed",
                status="error",
                agent_id=agent_id,
                kb_id=kb_id,
                source_type=source_type,
                error_code=error["code"],
                latency_ms=int((time.perf_counter() - started) * 1000),
            )
            logger.error("[kb] failed {}", redact_sensitive({"kbId": kb_id, "code": error["code"], "error": str(exc)}))
            result = {**result_base, **error}
            await _notify_progress(progress, result)
            results.append(result)

    return results


def _budget_for_agent(agent_id: str, payload: dict) -> dict[str, int]:
    budget = {
        "max_documents_per_job": MAX_DOCUMENTS_PER_JOB,
        "max_chunks_per_document": MAX_CHUNKS_PER_DOCUMENT,
    }

    env_budgets = _parse_budgets_json(os.environ.get("KB_AGENT_BUDGETS_JSON", ""))
    payload_budgets = payload.get("budgets") if isinstance(payload.get("budgets"), dict) else {}
    for source in (env_budgets, payload_budgets):
        candidate = source.get(agent_id) if isinstance(source.get(agent_id), dict) else source
        if not isinstance(candidate, dict):
            continue
        budget["max_documents_per_job"] = _positive_int(
            candidate,
            "maxDocumentsPerJob",
            "max_documents_per_job",
            default=budget["max_documents_per_job"],
        )
        budget["max_chunks_per_document"] = _positive_int(
            candidate,
            "maxChunksPerDocument",
            "max_chunks_per_document",
            default=budget["max_chunks_per_document"],
        )

    return budget


def _parse_budgets_json(value: str) -> dict:
    if not value:
        return {}
    try:
        parsed = json.loads(value)
    except Exception:
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _positive_int(source: dict, *keys: str, default: int) -> int:
    for key in keys:
        value = source.get(key)
        if value is None:
            continue
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            continue
        if parsed > 0:
            return parsed
    return default
